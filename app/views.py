from django.shortcuts import render, redirect
from . models import *
# Create your views here.\
import base64
from django.core.files.base import ContentFile
from django.contrib import messages
from tensorflow.keras.models import load_model
import imutils, pickle
import numpy as np
from tensorflow.keras.applications.mobilenet_v2 import preprocess_input
from tensorflow.keras.preprocessing.image import img_to_array
from sklearn.preprocessing import LabelEncoder
from django.db.models import Q
import cv2
import os
import uuid
from PIL import Image
import PyPDF2
import pytesseract
from django.http import JsonResponse
#from transformers import T5ForConditionalGeneration, T5Tokenizer
from django.core.paginator import Paginator

# Create your views here.
def index(request):
    return render(request, 'index.html')

def about(request):
    return render(request, 'about.html')


def register(request):
    if request.method == 'POST':
        name = request.POST.get('username')
        email = request.POST.get('email')
        password = request.POST.get('password')

        if UserModel.objects.filter(email=email).exists():
            messages.error(request, 'Email already registered!')
            return redirect('register')

        user = UserModel(username=name, email=email, password=password)
        user.save()
        messages.success(request, 'Registration successful! Please login.')
        return redirect('login')

    return render(request, 'register.html')

'''def training(request):
    le = LabelEncoder()
    faces, Id = getImagesAndLabels("UserImages")
    Id = le.fit_transform(Id)
    
    # Save encoder to a file
    with open('model/encoder.pkl', 'wb') as output:
        pickle.dump(le, output)
    
    # Train the recognizer
    recognizer = cv2.face.LBPHFaceRecognizer_create()
    recognizer.train(faces, np.array(Id))
    
    # Save the trained model
    recognizer.save(r"model/Trainner.yml")
    
    messages.success(request, 'Your model has been trained successfully!')
    return redirect('login')

def getImagesAndLabels(path):
    imagePaths = [os.path.join(path, f) for f in os.listdir(path)]
    faces = []
    Ids = []
    
    for imagePath in imagePaths:
        pilImage = Image.open(imagePath).convert('L')  # Convert to grayscale
        imageNp = np.array(pilImage, 'uint8')
        
        Id = str(os.path.split(imagePath)[-1].split(".")[0])  # Extract user ID
        faces.append(imageNp)
        Ids.append(Id)
    
    return faces, Ids

def detect_and_predict_person(img, faceNet, model):
    (h, w) = img.shape[:2]
    blob = cv2.dnn.blobFromImage(img, 1.0, (224, 224), (104.0, 177.0, 123.0))  # Pre-process image
    faceNet.setInput(blob)
    detections = faceNet.forward()  # Detect faces
    
    faces = []
    locs = []
    preds = []
    
    for i in range(detections.shape[2]):
        confidence = detections[0, 0, i, 2]
        if confidence > 0.5:  # Confidence threshold
            box = detections[0, 0, i, 3:7] * np.array([w, h, w, h])
            (startX, startY, endX, endY) = box.astype("int")
            (startX, startY) = (max(0, startX), max(0, startY))
            (endX, endY) = (min(w - 1, endX), min(h - 1, endY))
            face = img[startY:endY, startX:endX]
            face = cv2.cvtColor(face, cv2.COLOR_BGR2RGB)
            face = cv2.resize(face, (32, 32))
            face = img_to_array(face)
            face = preprocess_input(face)
            
            faces.append(face)
            locs.append((startX, startY, endX, endY))
    
    if len(faces) > 0:
        faces = np.array(faces, dtype="float32")
        preds = model.predict(faces, batch_size=32)
    
    return (locs, preds)'''

def login(request):
    if request.method == 'POST':
        email = request.POST['email']
        password = request.POST['password']

        if UserModel.objects.filter(email=email, password=password).exists():
            otp = str(random.randint(100000, 999999))  # 6-digit OTP
            request.session['otp'] = otp
            request.session['email'] = email

            # Send OTP via email
            send_mail(
                'Your 2FA OTP Code',
                f'Your OTP code is: {otp}',
                settings.DEFAULT_FROM_EMAIL,
                [email],
                fail_silently=False,
            )

            return redirect('verify_user_otp')  # A new page to enter OTP
        else:
            messages.error(request, 'Invalid email or password!')
            return redirect('login')

    return render(request, 'login.html')


import random
from django.core.mail import send_mail
def forgotpass(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        if UserModel.objects.filter(email=email).exists():
            user = UserModel.objects.get(email=email)
            
            otp = random.randint(10000,99999)
            user.otp = otp
            user.save()
            email_subject = 'Reset Passward Details'
            email_message = f'Hello {email},\n\nWelcome To Our Website!\n\nHere are your OTP details:\nEmail: {email}\OTP: {otp}\n\nPlease keep this information safe.\n\nBest regards,\nYour Website Team'
            from_email = 'takkellapativikram56@gmail.com'
            send_mail(email_subject, email_message, from_email, [email])
            messages.success(request, 'OTP sent successfully')
            return redirect('resetpassword')
        else:
            messages.error(request, 'Invalid email!')
            return redirect('forgotpass')


    return render(request, 'forgotpass.html')


def resetpassword(request):
    if request.method == 'POST':
        otp = request.POST.get('otp')
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirm_password')
        email = request.POST.get('email')
        if UserModel.objects.filter(otp=otp, email=email).exists():
            user = UserModel.objects.get(email=email)
            if password == confirm_password:
                user.password = password
                user.save()
                messages.success(request, 'Password reset successfully')
                return redirect('login')
            else:
                messages.error(request, 'Password and confirm password does not match!')
                return redirect('resetpassword')
        else:
            messages.error(request, 'Invalid OTP!')
            return redirect('resetpassword')
    return render(request, 'resetpassword.html')

def verify_user_otp(request):  # renamed to match URL
    if request.method == 'POST':
        entered_otp = request.POST['otp']
        original_otp = request.session.get('otp')

        if entered_otp == original_otp:
            request.session['is_authenticated'] = True
            return redirect('home')  # Change to your actual user dashboard
        else:
            messages.error(request, 'Invalid OTP!')
            return redirect('verify_user_otp')

    return render(request, 'verify_user_otp.html')

def resend_user_otp(request):
    """Resend OTP for regular users and update the session"""
    email = request.session.get("email")

    if not email:
        messages.error(request, "Session expired. Please start again.")
        return redirect("login")

    otp = send_otp(email)
    if otp:
        request.session['otp'] = otp  # Add this to store the new OTP
        messages.success(request, "A new OTP has been sent to your email.")
    else:
        messages.error(request, "Failed to send OTP. Please try again.")

    return redirect("verify_user_otp")


def home(request):
     email =request.session['email']
    #  print(email)
     return render(request, 'home.html',{'email':email})

import hashlib

def hash_string(input_string):
    # Create a SHA-256 hash object
    sha256_hash = hashlib.sha256()
    # Update the hash object with the bytes of the input string
    sha256_hash.update(input_string.encode('utf-8'))
    # Return the hexadecimal representation of the hash
    return sha256_hash.hexdigest()

import os
import uuid
from datetime import datetime
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import UploadFileModel, UserModel  # Import your models
import hashlib

def hash_string(content):
    """Generate a hash for the file content."""
    return hashlib.sha256(content.encode('utf-8')).hexdigest()

'''def uploadfiles(request):
    if request.method == 'POST':
        file = request.FILES['file']
        original_filename = file.name
        extension = os.path.splitext(original_filename)[1].lower()

        # Generate a unique filename with timestamp and UUID
        unique_filename = f"{os.path.splitext(original_filename)[0]}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}{extension}"

        file_path = os.path.join('static/assets', 'Files', unique_filename)
        os.makedirs(os.path.dirname(file_path), exist_ok=True)  # Ensure directory exists

        # Save the file
        with open(file_path, 'wb+') as f:
            for chunk in file.chunks():
                f.write(chunk)

        # Read file content safely
        if extension in ['.txt', '.csv', '.log']:  # Only read text files
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                messages.error(request, "Error decoding text file. Ensure it is UTF-8 encoded.")
                return redirect('uploadfiles')
        else:
            # For binary files like PDFs, DOCs, etc.
            with open(file_path, 'rb') as f:
                text = f.read()  # Read binary data
            text = str(hashlib.sha256(text).hexdigest())  # Generate hash for binary files

        # Hash and check for duplicates
        hashed_string = hash_string(text)
        if UploadFileModel.objects.filter(datahash=hashed_string).exists():
            messages.success(request, 'Data already exists!')
            return redirect('uploadfiles')

        user = UserModel.objects.get(email=request.session['email'])
        UploadFileModel.objects.create(
            filename=unique_filename,
            file=file_path,
            user=user,
            datahash=hashed_string
        ).save()

        messages.success(request, f"File '{original_filename}' uploaded successfully as '{unique_filename}'")
        return redirect('uploadfiles')

    return render(request, 'uploadfiles.html')'''


import os
import uuid
import hashlib
from datetime import datetime
import magic  # Detect file type (MIME)
import chardet  # Auto-detect file encoding
import fitz  # PyMuPDF for PDF text extraction
from docx import Document  # Extract text from DOCX files
import pythoncom  # Required for pywin32 on Windows
import win32com.client  # Extract text from DOC files (Microsoft Word)
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import UploadFileModel, UserModel

# Allowed file extensions
ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.csv'}
BLOCKED_EXTENSIONS = {'.mp3', '.mp4', '.zip', '.bin','.exe','.log','.jpg','.png','.py','.php','.java'}

# Function to detect encoding
def detect_encoding(file_path, default='utf-8'):
    with open(file_path, 'rb') as f:
        raw_data = f.read(10000)  # Read first 10,000 bytes for detection
    result = chardet.detect(raw_data)
    return result.get('encoding', default)

# Extract text from PDFs
def extract_text_from_pdf(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = "\n".join([page.get_text("text") for page in doc])
        return text.strip()
    except Exception:
        return None

# Extract text from DOCX files
def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs]).strip()
    except Exception:
        return None

# Extract text from DOC files using Microsoft Word
def extract_text_from_doc(doc_path):
    try:
        pythoncom.CoInitialize()  # Required for Windows
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # Run in background
        doc = word.Documents.Open(os.path.abspath(doc_path))
        text = doc.Content.Text.strip()
        doc.Close(False)
        word.Quit()
        return text
    except Exception:
        return None
def extract_text(file_path, extension):
    """Extract text from a file based on its extension."""
    if extension in {'.txt', '.csv'}:  # Text-based files
        encoding = detect_encoding(file_path)
        with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
            return f.read().strip()

    elif extension == '.pdf':  # PDFs
        return extract_text_from_pdf(file_path)

    elif extension == '.docx':  # Word DOCX
        return extract_text_from_docx(file_path)

    elif extension == '.doc':  # Word DOC (older format)
        return extract_text_from_doc(file_path)

    return None 
# Hash function to check duplicates
def hash_string(content):
    return hashlib.sha256(content.encode('utf-8')).hexdigest()



# 🔹 Upload File & Generate Insights
from django.shortcuts import render, redirect
from django.contrib import messages
import os
import uuid
import json
import hashlib

def extract_text(file_path, extension):
    """Extracts text from a document based on its file type."""
    # Implement extraction logic for PDFs, DOCX, TXT, etc.
    return None  # Placeholder

def hash_string(text):
    """Hashes the extracted text or file content."""
    return hashlib.sha256(text.encode()).hexdigest()

from sklearn.feature_extraction.text import CountVectorizer
from collections import Counter
import re

def generate_document_insights(text):
    # Simple keyword extraction
    words = re.findall(r'\w+', text.lower())
    common_words = Counter(words).most_common(10)
    keywords = [word for word, _ in common_words]

    # Basic summarization (just first few lines for now)
    summary = "\n".join(text.strip().splitlines()[:5])

    # Basic category (you can improve with ML later)
    if "invoice" in text.lower():
        category = "Finance"
    elif "project" in text.lower():
        category = "Project Report"
    else:
        category = "General Document"

    return {
        "summary": summary,
        "keywords": keywords,
        "category": category
    }


import os
import uuid
import json
import hashlib
from datetime import datetime
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import UploadFileModel, UserModel
#from .utils import extract_text, generate_document_insights, hash_string  # adjust if your helpers are in a different place

  # Update with your own blocked types
def calculate_similarity(hash1, hash2):
    if not hash1 or not hash2:
        return 0.0  # or handle accordingly
    matches = sum(1 for a, b in zip(hash1, hash2) if a == b)
    return matches / max(len(hash1), len(hash2))  # Use max to avoid division by 0


def uploadfiles(request):
    if request.method == 'POST':
        if 'file' not in request.FILES:
            messages.error(request, "No file selected for upload.")
            return redirect('uploadfiles')

        file = request.FILES['file']
        original_filename = file.name
        extension = os.path.splitext(original_filename)[1].lower()

        if extension in BLOCKED_EXTENSIONS:
            messages.error(request, "Invalid file type! Only allowed file types.")
            return redirect('uploadfiles')

        # Generate unique filename and save path
        unique_filename = f"{os.path.splitext(original_filename)[0]}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}{extension}"
        file_dir = os.path.join('static/assets', 'Files')
        os.makedirs(file_dir, exist_ok=True)
        file_path = os.path.join(file_dir, unique_filename)

        # Save file to disk
        with open(file_path, 'wb+') as f:
            for chunk in file.chunks():
                f.write(chunk)

        # Extract text and generate insights
        text = extract_text(file_path, extension)
        if text is None:
            with open(file_path, 'rb') as f:
                text = hashlib.sha256(f.read()).hexdigest()

        try:
            insights = generate_document_insights(text)
        except Exception as e:
            print(f"[Insight Error]: {e}")
            insights = {
                "summary": "Failed to generate insights.",
                "keywords": [],
                "category": "Unknown"
            }

        # Save insights to JSON
        insights_path = file_path.replace(extension, ".json")
        try:
            with open(insights_path, 'w', encoding='utf-8') as f:
                json.dump(insights, f, indent=4)
            print(f"[INFO] Insights saved at: {insights_path}")
        except Exception as e:
            print(f"[ERROR] Failed to save insights: {e}")
            messages.error(request, "Failed to generate insights.")
            return redirect('uploadfiles')

        # Hash the file/text content
        hashed_string = hash_string(text)

        # Get current user
        user = UserModel.objects.get(email=request.session['email'])

        # ✅ Check if file with same hash is already uploaded by other user
        duplicate_file = UploadFileModel.objects.filter(datahash=hashed_string).exclude(user=user).first()
        if duplicate_file:
            UploadFileModel.objects.create(
                filename=unique_filename,
                file=file_path,
                user=user,
                datahash=hashed_string,
                is_duplicate=True,
                status='pending',
                insights_path=insights_path
            )
            messages.warning(request, f"This file has already been uploaded by user: {duplicate_file.user.username}. "
                                      f"Your file is under admin review.")
            return redirect('uploadfiles')

        # ✅ Check if the current user already uploaded the same file
        existing_file = UploadFileModel.objects.filter(datahash=hashed_string, user=user).first()
        if existing_file:
            messages.warning(request, "You have already uploaded this file.")
            return redirect('uploadfiles')

        # ✅ Save as a new unique file
        UploadFileModel.objects.create(
            filename=unique_filename,
            file=file_path,
            user=user,
            datahash=hashed_string,
            is_duplicate=False,
            status='approved',
            insights_path=insights_path
        )

        messages.success(request, f"File '{original_filename}' uploaded successfully with AI insights.")
        return redirect('uploadfiles')

    return render(request, 'uploadfiles.html')


from django.shortcuts import get_object_or_404
import os

def insight(request, file_id):
    email = request.session.get('email')
    user = UserModel.objects.get(email=email)

    # Get only the file uploaded by the current user with matching ID
    file = get_object_or_404(UploadFileModel, id=file_id, user=user)

    insights = {
        "filename": file.filename,
        "summary": "No insights available for this document.",
        "keywords": [],
        "category": "Uncategorized"
    }
    info_path = None
    # Load insights from file if available
    if file.insights_path and os.path.exists(file.insights_path):
        try:
            with open(file.insights_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
                insights["summary"] = json_data.get("summary", "No summary available")
                insights["keywords"] = json_data.get("keywords", [])
                insights["category"] = json_data.get("category", "General Document")

                info_path = f"[INFO] Insights saved at: {file.insights_path}"
        except Exception as e:
            print(f"[ERROR] Error loading insights for {file.filename}: {e}")
            messages.error(request, "Failed to load insights.")
    else:
        print(f"[ERROR] Insights file not found for {file.filename}")

    return render(request, 'insight.html', {
        'email': email,
        'insights': insights,
        'file_name': file.filename,
        'info_path': info_path
    })


from django.core.paginator import Paginator
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from .models import UploadFileModel, UserModel

  # Ensure user is logged in
def viewfiles(request):
    email = request.session.get("email")  # Get email from session
    
    if not email:
        return redirect("login")  # Redirect if no email in session

    try:
        user = UserModel.objects.get(email=email)
    except UserModel.DoesNotExist:
        return redirect("login")

    # Fetch only approved files for the user
    files = UploadFileModel.objects.filter(user=user, is_duplicate=False).exclude(status="Rejected").order_by("-uploaded_at")

    # Pagination (4 files per page)
    paginator = Paginator(files, 4)
    page_number = request.GET.get("page")
    page_data = paginator.get_page(page_number)

    return render(request, "viewfiles.html", {"data": page_data, "email": email})

from django.shortcuts import render
from gtts import gTTS
import os
from datetime import datetime
import uuid
import chardet  # Auto-detect file encoding
import fitz  # PyMuPDF for PDF text extraction
from docx import Document  # Extract text from DOCX files
import pythoncom  # Required for pywin32 on Windows
import win32com.client  # Extract text from DOC files (Microsoft Word)
from .models import UploadFileModel

# Supported file types
TEXT_FILE_EXTENSIONS = {'.txt'}
PDF_FILE_EXTENSIONS = {'.pdf'}
DOCX_FILE_EXTENSIONS = {'.docx'}
DOC_FILE_EXTENSIONS = {'.doc'}

def detect_encoding(file_path, default='utf-8'):
    """Detect file encoding to avoid decoding errors."""
    with open(file_path, 'rb') as f:
        raw_data = f.read(10000)  # Read first 10,000 bytes for detection
    result = chardet.detect(raw_data)
    return result.get('encoding', default)

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(pdf_path)  # Open the PDF file
        text = ""
        for page in doc:
            text += page.get_text("text")  # Extract text from each page
        return text.strip()
    except Exception as e:
        return None  # Return None if extraction fails

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = Document(docx_path)  # Open the DOCX file
        text = "\n".join([para.text for para in doc.paragraphs])  # Extract paragraphs as text
        return text.strip()
    except Exception as e:
        return None  # Return None if extraction fails

def extract_text_from_doc(doc_path):
    """Extract text from a DOC file using pywin32 (Microsoft Word)."""
    try:
        pythoncom.CoInitialize()  # Initialize COM (Required for Windows)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # Run MS Word in the background

        doc = word.Documents.Open(os.path.abspath(doc_path))  # Open the document
        text = doc.Content.Text.strip()  # Extract text
        doc.Close(False)  # Close the document (without saving)
        word.Quit()  # Quit Word application

        return text
    except Exception as e:
        return None  # Return None if extraction fails

def text_to_speech(request, id, req):
    # Retrieve the file path from the model
    if req == 'owner':
        # print('skfsdnfd',req)
        data = UploadFileModel.objects.get(id=id)
        file_path = data.file.path
    else:
        # print('ppdddkd',req)
        data=RequestFileModel.objects.get(id=id)
        file_path = data.file_id.file.path
          
    extension = os.path.splitext(file_path)[1].lower()  # Get file extension

    text = ""

    if extension in TEXT_FILE_EXTENSIONS:
        # Detect file encoding
        encoding = detect_encoding(file_path)
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
        except (UnicodeDecodeError, IOError):
            return render(request, 'text_to_speech.html', {'error': "Unable to read file content. Ensure it's a valid text file."})

    elif extension in PDF_FILE_EXTENSIONS:
        text = extract_text_from_pdf(file_path)
        if not text:
            return render(request, 'text_to_speech.html', {'error': "Unable to extract text from PDF file."})

    elif extension in DOCX_FILE_EXTENSIONS:
        text = extract_text_from_docx(file_path)
        if not text:
            return render(request, 'text_to_speech.html', {'error': "Unable to extract text from DOCX file."})

    elif extension in DOC_FILE_EXTENSIONS:
        text = extract_text_from_doc(file_path)
        if not text:
            return render(request, 'text_to_speech.html', {'error': "Unable to extract text from DOC file."})

    else:
        return render(request, 'text_to_speech.html', {'error': "Unsupported file format. Please upload a valid text, PDF, or DOC file."})

    # Convert extracted text to speech
    tts = gTTS(text=text, lang='en')
    filename = "speech.mp3"
    extension = os.path.splitext(filename)[1]
    unique_filename = f"{os.path.splitext(filename)[0]}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}{extension}"
    output_file_path = os.path.join('static/assets/audio', unique_filename)

    os.makedirs(os.path.dirname(output_file_path), exist_ok=True)  # Ensure directory exists
    tts.save(output_file_path)  # Save the audio file

    return render(request, 'text_to_speech.html', {'audio_file': output_file_path})


def logout(request):
    del request.session['email']
    return redirect('index')


def profile(request):
    # UserProfile.objects.all().delete()
    email =  request.session['email']
    user = UserModel.objects.get(email=email)
    data = UserProfile.objects.filter(user_id=user.id).exists()
    if data:
        profile = UserProfile.objects.filter(user_id=user.id)

        user = UserModel.objects.filter(email=email)
        return render(request, 'profile.html', {'profile': profile,'user': user})
    else:
        userdata = UserModel.objects.filter(email=email)

        return render(request, 'updateprofile.html',{'user':userdata})
    
    # return render(request, 'profile.html')

def updateprofile(request):
    email =  request.session['email']
    user = UserModel.objects.get(email=email)
    if request.method == 'POST':
        # name = request.POST.get('name')
        phone = request.POST.get('phone')
        address = request.POST.get('location')
        bio = request.POST['bio']
        image = request.FILES['image']
        data = UserProfile.objects.create(
            user_id=user.id,
           
            phone=phone,
            address=address,
            bio=bio,
            image=image 
        
        )
        data.save()
        return redirect('profile')

def editprofile(request):
    # UserProfile.objects.filter(id=4).delete()
    email =  request.session['email']
    user = UserModel.objects.get(email=email)
    profile = UserModel.objects.filter(email=email)
    if request.method == 'POST':
        phone = request.POST.get('phone')
        address = request.POST.get('location')
        bio = request.POST['bio']
        image = request.FILES['image']
        data = UserProfile.objects.get(user_id=user.id)
        if phone:
            data.phone = phone
        if address :
            data.address = address
        data.bio = bio
        if image:
            data.image = image
        data.save()
        messages.success(request, 'Profile Updated Successfully!')
        return redirect('profile')
    return render(request, 'editprofile.html',{'user':profile})


def sendrequest(request, id):
    requester = request.session['email']
    data = UploadFileModel.objects.get(id=id)
    if RequestFileModel.objects.filter(requester=requester, file_id=data).exists():
        messages.success(request, 'You already requested this file')
        return redirect('viewfiles')
    else:
        req = RequestFileModel.objects.create(
            file_id = data,
            requester = requester
        )
        req.save()
        messages.success(request, 'Request Sent Successfully!')
        return redirect('viewfiles')

        

def viewrequests(request):
    email = request.session['email']
    user = UserModel.objects.get(email=email)  # Fetch the user object

    requests = RequestFileModel.objects.filter(file_id__user__email=email)
    paginator = Paginator(requests, 4)
    page_number = request.GET.get('page')
    page_data = paginator.get_page(page_number)

    return render(request, 'viewrequests.html', {
        'data': page_data,
        'email': email,
        
    })



def acceptrequest(request, id):
    email = request.session['email']
    req = RequestFileModel.objects.get(id=id)
    req.status = 'Accepted'
    req.save()
    messages.success(request, 'Request Accepted Successfully!')
    return redirect('viewrequests')


def viewresponses(request):
    email =  request.session['email']
    responses = RequestFileModel.objects.filter(requester=email, status='Accepted')
    paginator = Paginator(responses, 4)
    page_number = request.GET.get('page')
    page_data = paginator.get_page(page_number)
    return render(request, 'viewresponses.html',{'data':page_data, 'email':email})

'''def insight(request):
    email =  request.session['email']
    return render(request, 'insight.html',{ 'email':email})

def insight(request, file_id):
    email = request.session.get('email', 'Guest')

    # Fetch the uploaded file details
    uploaded_file = UploadFileModel.objects.get(id=file_id)
    
    # Load insights from JSON file
    insights_path = uploaded_file.file.replace(os.path.splitext(uploaded_file.file)[1], ".json")
    insights = {}
    if os.path.exists(insights_path):
        with open(insights_path, 'r', encoding='utf-8') as f:
            insights = json.load(f)

    context = {
        'email': email,
        'file': uploaded_file.filename,
        'insights': insights
    }
    return render(request, 'insight.html', context)
'''

import random
from django.shortcuts import render, redirect
from django.contrib import messages
from django.core.mail import send_mail
from .models import AdminUserModel

def admin_register(request):
    if request.method == 'POST':
        admin_id = request.POST.get('admin_id')
        email = request.POST.get('email')
        password = request.POST.get('password')

        if not admin_id.startswith('DV'):
            messages.error(request, "Admin ID must start with 'DV'.")
            return redirect('admin_register')

        if AdminUserModel.objects.filter(email=email).exists():
            messages.error(request, "Email already registered!")
            return redirect('admin_register')

        otp = random.randint(10000, 99999)

        admin_user = AdminUserModel(admin_id=admin_id, email=email, password=password, otp=otp)
        admin_user.save()

        # Send OTP via email
        email_subject = "Verify Your Admin Account"
        email_message = f"Your OTP for Admin Verification: {otp}"
        send_mail(email_subject, email_message, "takkellapativikram56@gmail.com", [email])

        messages.success(request, "OTP sent! Verify your email.")
        return redirect('verify_admin_otp', admin_id=admin_id)

    return render(request, 'admin_register.html')
import random
from django.shortcuts import render, redirect
from django.core.mail import send_mail
from django.contrib import messages
from django.contrib.auth.models import User

# Simulate an OTP storage (should use session or database)
OTP_STORAGE = {}

def generate_otp():
    """Generate a 6-digit OTP"""
    return str(random.randint(100000, 999999))

def send_otp_email(email, otp):
    """Send OTP via email"""
    subject = "Your OTP Verification Code"
    message = f"Your OTP code is {otp}. It is valid for 5 minutes."
    from_email = "takkellapativikram56@gmail.com"  # Replace with your email
    recipient_list = [email]

    send_mail(subject, message, from_email, recipient_list)

def verify_admin_otp(request):
    """Verify OTP entered by the admin"""
    if request.method == "POST":
        entered_otp = request.POST.get("otp")
        stored_otp = request.session.get("admin_otp")

        if entered_otp and str(entered_otp) == str(stored_otp):
            del request.session['admin_otp']  # Remove OTP after successful login
            return redirect('admin_dashboard')
        else:
            messages.error(request, "Invalid OTP! Please try again.")

    return render(request, "verify_admin_otp.html")

import random
from django.core.mail import send_mail
from django.shortcuts import render, redirect
from django.contrib import messages
from django.conf import settings
from .models import AdminUserModel

# OTP Storage (Temporary Dictionary)
OTP_STORAGE = {}

def generate_otp():
    """Generate a 6-digit OTP"""
    return str(random.randint(100000, 999999))

def send_otp(email):
    """Generate and send OTP via email"""
    otp = generate_otp()
    subject = "Admin Login OTP Verification"
    message = f"Your OTP for login is: {otp}. Do not share it with anyone."

    try:
        send_mail(subject, message, settings.EMAIL_HOST_USER, [email])
        OTP_STORAGE[email] = otp  # Store OTP in memory (Use DB in production)
        return otp
    except Exception as e:
        print("Error sending OTP:", e)
        return None

def resend_otp(request):
    """Resend OTP and update the session"""
    email = request.session.get("admin_email")  # Ensure correct session key

    if not email:
        messages.error(request, "Session expired. Please start over.")
        return redirect("login")  

    otp = send_otp(email)  # Send new OTP
    if otp:
        messages.success(request, "A new OTP has been sent to your email.")
    else:
        messages.error(request, "Failed to send OTP. Try again!")

    return redirect("verify_admin_otp")  # Redirect back to verification page

from django.shortcuts import render
import hashlib
from .models import UploadFileModel, AdminUserModel

def check_duplicate(request):
    if request.method == "POST":
        uploaded_file = request.FILES["file"]
        file_hash = hashlib.md5(uploaded_file.read()).hexdigest()

        existing_files = UploadFileModel.objects.all()
        for file in existing_files:
            similarity = calculate_similarity(file.file_hash, file_hash)
            print(similarity)

            if similarity >= 89:
                file.is_duplicate = True
                file.save()
                return render(request, "uploadfiles.html", 
                {"file": file,
                "similarity": similarity})

        UploadFileModel.objects.create(file=uploaded_file, file_hash=file_hash)
        return redirect("viewfiles")

    return render(request, "uploadfiles.html")

def admin_verify_files(request):
    if "admin_id" not in request.session:
        return redirect("admin_login")

    flagged_files = UploadFileModel.objects.filter(is_duplicate=True)
    return render(request, "admin_verify_files.html", {"flagged_files": flagged_files})


def admin_approve_document(request, id):
    if "admin_id" not in request.session:
        return redirect("admin_login")

    file = UploadFileModel.objects.get(id=id)
    file.is_duplicate = False
    file.save()
    messages.success(request, "File Approved!")
    return redirect("admin_verify_files")
def admin_dashboard(request):
    return render(request,"admin_dashboard.html")
def admin_view_files(request):
    files = UploadFileModel.objects.all()
    
    # Paginate the files (4 per page)
    paginator = Paginator(files, 4)
    page_number = request.GET.get('page')
    page_data = paginator.get_page(page_number)

    return render(request, 'admin_view_files.html', {'data': page_data})

from django.shortcuts import render, redirect
from django.contrib import messages
from django.core.mail import send_mail
from django.conf import settings
from .models import AdminUserModel

# OTP Storage (Temporary Dictionary)
OTP_STORAGE = {}

def generate_otp():
    """Generate a 6-digit OTP"""
    return str(random.randint(100000, 999999))

def send_otp(email):
    """Generate and send OTP via email"""
    otp = generate_otp()
    subject = "Admin Login OTP Verification"
    message = f"Your OTP for login is: {otp}. Do not share it with anyone."

    try:
        send_mail(subject, message, settings.EMAIL_HOST_USER, [email])
        OTP_STORAGE[email] = otp  # Store OTP in memory (Use DB in production)
        return otp
    except Exception as e:
        print("Error sending OTP:", e)
        return None

def resend_otp(request):
    """Resend OTP and update the session"""
    email = request.session.get("admin_email")  # Ensure correct session key

    if not email:
        messages.error(request, "Session expired. Please start over.")
        return redirect("login")  

    otp = send_otp(email)  # Send new OTP
    if otp:
        messages.success(request, "A new OTP has been sent to your email.")
    else:
        messages.error(request, "Failed to send OTP. Try again!")

    return redirect("verify_admin_otp")  # Redirect back to verification page

def admina_login(request):
    if request.method == "POST":
        admin_id = request.POST.get("admin_id")
        password = request.POST.get("password")

        try:
            admin = AdminUserModel.objects.get(admin_id=admin_id)
        except AdminUserModel.DoesNotExist:
            messages.error(request, "Invalid Admin ID!")
            return redirect('admina_login')

        # Check password
        if admin.password != password:
            messages.error(request, "Incorrect password!")
            return redirect('admina_login')

        # Generate new OTP for every login attempt
        otp = send_otp(admin.email)
        if otp:
            request.session['admin_otp'] = otp  # Store new OTP
            request.session['admin_email'] = admin.email  # Store email for verification
            messages.success(request, "A new OTP has been sent to your email.")
            return redirect('verify_admin_otp')  # Redirect to OTP verification
        else:
            messages.error(request, "Failed to send OTP. Try again!")

    return render(request, "admin_login.html")

def verify_admin_otp(request):
    """Verify OTP entered by the admin"""
    if request.method == "POST":
        entered_otp = request.POST.get("otp")
        stored_otp = request.session.get("admin_otp")

        if entered_otp and str(entered_otp) == str(stored_otp):
            del request.session['admin_otp']  # Remove OTP after successful login
            return redirect('admin_dashboard')
        else:
            messages.error(request, "Invalid OTP! Please try again.")

    return render(request, "verify_admin_otp.html")

def admin_dashboard(request):
    files = UploadFileModel.objects.all()
    
    # Paginate the files (4 per page)
    paginator = Paginator(files, 4)
    page_number = request.GET.get('page')
    page_data = paginator.get_page(page_number)

    return render(request, 'admin_view_files.html', {'data': page_data})


def pending_approvals(request):
    pending_admins = AdminUserModel.objects.filter(is_verified=False)
    uploaded_files = UploadFileModel.objects.all()

    file_hashes = {}
    duplicate_files = []

    for file in uploaded_files:
        if file.datahash in file_hashes:
            similarity = calculate_similarity(file.datahash, file_hashes[file.datahash].datahash)
            if similarity > 90:  # Threshold to consider as duplicate
                file.is_duplicate = True
                file.similarity = similarity
                file.save()
                duplicate_files.append(file)
        else:
            file_hashes[file.datahash] = file

    context = {
        "pending_admins": pending_admins,
        "uploaded_files": uploaded_files,
        "duplicate_files": duplicate_files,
    }
    return render(request, "pending_approvals.html", context)


    # Identify duplicate files (by hash, name, etc.)
    '''from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import AdminUserModel, UploadFileModel  # Import models

def pending_approvals(request):
    # Fetch all pending admin approvals
    pending_admins = AdminUserModel.objects.filter(is_verified=False)

    # Fetch all uploaded files by users
    uploaded_files = UploadFileModel
    file_hashes = {}
    duplicate_files = []

    for file in uploaded_files:
        if file.file_hash in file_hashes:
            duplicate_files.append(file)
        else:
            file_hashes[file.file_hash] = file

    context = {
        "pending_admins": pending_admins,
        "uploaded_files": uploaded_files,
        "duplicate_files": duplicate_files,  # Pass duplicates separately
    }
    return render(request, "pending_approvals.html", context)'''
from django.shortcuts import render, get_object_or_404, redirect
def accept_file(request, file_id):
    file = get_object_or_404(UploadFileModel, id=file_id)
    file.status = "Accepted"
    file.is_duplicate = False 
    file.save()
    messages.success(request, "File has been accepted.")
    return redirect("admin_view_files")

def reject_file(request, file_id):
    file = get_object_or_404(UploadFileModel, id=file_id)
    file_entry = get_object_or_404(UploadFileModel, id=file_id)
    file.status = "Rejected"
    file.save()
    subject = "File Rejection Notification"
    message = (
        f"Hello {file_entry.user.username},\n\n"
        f"Your uploaded file '{file_entry.filename}' has been rejected by the admin.\n"
        f"Reason: Please contact support for more details.\n\n"
        f"Regards,\nAdmin Team"
    )

    send_mail(
        subject,
        message,
        settings.DEFAULT_FROM_EMAIL,  
        [file_entry.user.email],  # Send email to the file uploader
        fail_silently=False,
    )

    messages.error(request, "File has been rejected and the user has been notified.")
    return redirect("admin_view_files")

def pending_file(request, file_id):
    file = get_object_or_404(UploadFileModel, id=file_id)
    file.status = "Pending"
    file.save()
    messages.info(request, "File is now pending review.")
    return redirect("admin_view_files")
from django.shortcuts import render
from .models import UploadFileModel  # Import your model if needed

def admin_view_files(request):
    data = UploadFileModel.objects.all()  # Fetch all files from the database
    return render(request, 'admin_file_verification.html', {'data': data})
def admin_logout(request):
    request.session.pop('email', None)  # This avoids KeyError if 'email' is missing
    return redirect('index')
def admin_approve_document(request, id):
    if "admin_id" not in request.session:
        return redirect("admin_login")

    file = UploadFileModel.objects.get(id=id)
    file.is_duplicate = False
    file.save()
    messages.success(request, "File Approved!")
    return redirect("admin_verify_files")

def admin_reject_document(request, id):
    if "admin_id" not in request.session:
        return redirect("admin_login")

    file = UploadFileModel.objects.get(id=id)
    file.delete()
    messages.error(request, "File Rejected!")
    return redirect("admin_verify_files")

def admin_set_pending(request, id):
    if "admin_id" not in request.session:
        return redirect("admin_login")

    file = UploadFileModel.objects.get(id=id)
    file.is_duplicate = True
    file.save()
    messages.info(request, "File set to Pending for further review!")
    return redirect("admin_verify_files")


from django.core.mail import send_mail
from django.conf import settings
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from .models import UploadFileModel

def notify_users(request, file_id):
    file_entry = get_object_or_404(UploadFileModel, id=file_id)
    
    # Find duplicate file using 'datahash' instead of 'file_hash'
    duplicate_file = UploadFileModel.objects.filter(datahash=file_entry.datahash).exclude(id=file_id).first()
    
    if duplicate_file:
        user1_email = file_entry.user.email
        user2_email = duplicate_file.user.email
        
        subject = "Duplicate File Uploaded - Admin Notification"
        message = (
            f"Hello,\n\n"
            f"The file you uploaded ({file_entry.filename}) is similar to another file uploaded by {duplicate_file.user.username}.\n"
            f"Please verify the data and contact support if needed or send the document to above email for furthur verification.\n\n"
            f"Regards,\nAdmin Team"
        )
        
        send_mail(
            subject,
            message,
            settings.DEFAULT_FROM_EMAIL,  
            [user1_email, user2_email],  # Sending mail to both users
            fail_silently=False,
        )
        
        messages.success(request, "Notification sent to both users.")
    else:
        messages.warning(request, "No duplicate found for this file.")

    return redirect("admin_dashboard")